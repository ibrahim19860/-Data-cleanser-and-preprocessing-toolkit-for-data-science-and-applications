# Re-run the document generation code since execution state was reset

from docx import Document

# Create a new Word document
doc = Document()

# Set title
doc.add_heading('Data Cleaning Toolkit: Results and Comparative Analysis', level=1)

# Add Results and Discussion section
doc.add_heading('Results and Discussion', level=2)

doc.add_paragraph(
    "The Data Cleaning Toolkit is designed to preprocess raw datasets containing various "
    "inconsistencies such as missing values, outliers, duplicate records, and inconsistent categorical "
    "variables. These inconsistencies can significantly impact machine learning model performance, "
    "leading to biased predictions and reduced accuracy. The toolkit automates the data cleaning "
    "process by applying statistical imputation, outlier detection, duplicate removal, categorical encoding, "
    "and feature scaling, ensuring that datasets are transformed into a structured and optimized format "
    "for machine learning applications."
)

doc.add_paragraph(
    "To evaluate the effectiveness of the toolkit, the raw dataset underwent a series of preprocessing steps. "
    "Missing values were handled using mean, median, or mode imputation, depending on the feature type. "
    "Outliers were detected and removed using Interquartile Range (IQR) and Z-score techniques. Duplicate "
    "entries were identified and eliminated to ensure dataset uniqueness. Categorical variables were standardized "
    "using one-hot encoding and label encoding, and numerical features were normalized and standardized to "
    "maintain consistency in data distribution."
)

# Add Table 1: Data Quality Before and After Cleaning
doc.add_heading('Table 1: Data Quality Before and After Cleaning', level=3)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Aspect'
hdr_cells[1].text = 'Before Cleaning'
hdr_cells[2].text = 'After Cleaning'

data = [
    ["Missing Values", "Data contained several missing values", "Imputed using mean, median, or mode"],
    ["Outliers", "Extreme values skewing feature distributions", "Outliers removed using IQR/Z-Score techniques"],
    ["Duplicates", "Redundant records increasing bias", "Duplicates removed for data integrity"],
    ["Categorical Data", "Inconsistent labels and missing values", "Encoded uniformly with structured categories"],
    ["Feature Scaling", "Some features had large variations in scale", "Normalized and standardized for uniformity"]
]

for row_data in data:
    row_cells = table.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

# Add Performance Evaluation Section
doc.add_heading('Performance Evaluation Using Machine Learning Models', level=2)

doc.add_paragraph(
    "To validate the impact of data cleaning on predictive performance, the cleaned dataset was used to "
    "train Random Forest, Logistic Regression, K-Nearest Neighbors (KNN), and Support Vector Machine (SVM) "
    "models. The models were trained and tested on both the raw dataset (before cleaning) and the preprocessed "
    "dataset (after cleaning). The accuracy results before and after data cleaning are presented in Table 2."
)

# Add Table 2: Model Accuracy Before and After Data Cleaning
doc.add_heading('Table 2: Model Accuracy Before and After Data Cleaning', level=3)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Machine Learning Model'
hdr_cells[1].text = 'Accuracy Before Cleaning'
hdr_cells[2].text = 'Accuracy After Cleaning'

accuracy_data = [
    ["Random Forest", "0.82", "0.96"],
    ["Logistic Regression", "0.78", "0.96"],
    ["KNN", "0.80", "0.98"],
    ["SVM", "0.79", "0.97"]
]

for row_data in accuracy_data:
    row_cells = table2.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

doc.add_paragraph(
    "The results demonstrate a significant improvement in model performance after applying data cleaning techniques. "
    "The KNN model showed the highest improvement, with accuracy increasing from 0.80 to 0.98, highlighting the "
    "impact of removing noisy data. SVM, Logistic Regression, and Random Forest models also exhibited notable accuracy "
    "improvements due to the elimination of outliers, proper handling of missing values, and feature standardization. "
    "These findings confirm that high-quality data directly influences model reliability and predictive accuracy."
)

# Save document
file_path = "C:\86"
doc.save(file_path)

# Provide download link
file_path
